"""
Report generation module.

This module handles the generation of comprehensive backtesting reports
including Word documents, charts, and CSV summaries.
"""

import os
import sys
import zipfile
import shutil
import tempfile
import tarfile
import json
import importlib.util
from pathlib import Path
from datetime import datetime
import logging
from io import BytesIO

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.figure
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from ..backtesting.engine import BacktestResult

logger = logging.getLogger(__name__)


class DocxLogger:
    """Custom logger that writes to a Word document."""
    
    def __init__(self, filepath):
        """Initialize the DocxLogger."""
        self.filepath = filepath
        self.document = Document()
        
        try:
            # Load existing document if it exists
            self.document = Document(filepath)
        except:
            pass  # Start a new document if file doesn't exist
        
        # Default styling
        self.font_name = "Times New Roman"
        self.font_size = 12
        self.space_after = Pt(12)
    
    def set_default_style(self, font_name="Times New Roman", 
                         font_size=12, space_after=12):
        """Set the default style for paragraphs."""
        self.font_name = font_name
        self.font_size = font_size
        self.space_after = Pt(space_after)
    
    def write(self, message):
        """Write a message to the document."""
        if message.strip():  # Avoid writing empty lines
            paragraph = self.document.add_paragraph(message)
            run = paragraph.runs[0]
            
            # Set fonts for both Western and East Asian text
            run.font.name = self.font_name
            run.font._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            run.font.size = Pt(self.font_size)
            
            # Set paragraph spacing
            paragraph.paragraph_format.space_after = self.space_after
    
    def write_empty_line(self):
        """Add an empty line to the document."""
        paragraph = self.document.add_paragraph("")
        paragraph.paragraph_format.space_after = self.space_after
    
    def add_plot(self, fig, plot_title=None):
        """Add a plot to the document."""
        # Save the plot as an image
        image_path = "temp_plot.png"
        fig.savefig(image_path, dpi=300, bbox_inches="tight")
        
        # Add title if provided
        if plot_title:
            title = self.document.add_heading(plot_title, level=2)
            title_format = title.paragraph_format
            title_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add the plot image with smaller width to fit in document margins
        self.document.add_picture(image_path, width=Inches(6.0))
        
        # Clean up temp file
        try:
            os.remove(image_path)
        except:
            pass
    
    def flush(self):
        """Save changes to the file."""
        self.document.save(self.filepath)


class ReportGenerator:
    """Generates comprehensive backtesting reports."""
    
    def __init__(self, strategies_directory="./oosit_strategies", config=None):
        """
        Initialize the ReportGenerator.
        
        Args:
            strategies_directory: Directory containing strategy ZIP files
            config: Configuration dictionary containing font settings
        """
        self.strategies_directory = Path(strategies_directory)
        self.temp_docx_dir = None
        self.backtest_results = None  # Store backtest results for default strategy access
        self.timestamp_dir = None
        self.csv_dir = None
        self.docx_dir = None
        self.archive_dir = None
        self.config = config or {}
    
    def generate_all_reports(self, backtest_results, 
                           config):
        """
        Generate all reports including Word documents and CSV summaries.
        
        Args:
            backtest_results: Results from BacktestEngine.run_full_backtest()
            config: Configuration dictionary
        """
        logger.info("Starting report generation")
        
        # Clean up existing result directories and create fresh ones
        self._setup_result_directories()
        
        # Store backtest results for default strategy access
        self.backtest_results = backtest_results
        
        try:
            # Extract DOCX files from strategy ZIPs
            self._extract_docx_files()
            
            # Generate individual strategy reports
            self._generate_strategy_reports(backtest_results, config)
            
            # Generate CSV summaries
            self._generate_csv_summaries(backtest_results, config)
            
            # Generate archive
            self._generate_archive(config)
            
            logger.info("Report generation completed successfully")
            
        except Exception as e:
            logger.error(f"Error generating reports: {e}")
            raise
        finally:
            self._cleanup_temp_directories()
    
    def _setup_result_directories(self):
        """Create directory structure with new naming convention."""
        # Use output_directory from config, default to './oosit_results' if not provided
        output_directory = self.config.get('output_directory', './oosit_results')
        results_root = Path(output_directory)
        results_root.mkdir(exist_ok=True)
        
        # Generate directory name based on test strategies and config file
        test_strategies = self.config.get('testing_strategies', [])
        config_file = self.config.get('config_file', None)
        
        # Create comma-joined string of test strategies
        strategies_str = ','.join(test_strategies) if test_strategies else 'no_test_strategies'
        
        # Determine additional flag
        additional_flag = ''
        if config_file:
            config_path = Path(config_file)
            config_name = config_path.stem  # Get filename without extension
            
            if config_name == 'default_config':
                additional_flag = ''
            elif config_name == 'dxy_config':
                additional_flag = ' [DXY]'
            elif config_name == 'veu_redirect_config':
                additional_flag = ' [VEU]'
            else:
                additional_flag = f' [{config_name}]'
        
        # Generate timestamp
        timestamp = datetime.now().strftime("%y%m%d-%H%M%S")
        
        # Create directory name
        dir_name = f"{strategies_str}{additional_flag} ({timestamp})"
        self.timestamp_dir = results_root / dir_name
        
        # Create main timestamp directory
        self.timestamp_dir.mkdir(exist_ok=True)
        logger.info(f"Created timestamp directory: {self.timestamp_dir}")
        
        # Create subdirectories for CSV and DOCX results
        self.csv_dir = self.timestamp_dir / 'csv results'
        self.docx_dir = self.timestamp_dir / 'docx results'
        self.archive_dir = self.timestamp_dir / 'run_archive'
        
        self.csv_dir.mkdir(exist_ok=True)
        self.docx_dir.mkdir(exist_ok=True)
        self.archive_dir.mkdir(exist_ok=True)
        
        logger.info(f"Created subdirectories: csv results, docx results, run_archive")
    
    def _extract_docx_files(self):
        """Create DOCX files for each strategy based on _explanation from .py files."""
        self.temp_docx_dir = tempfile.mkdtemp(prefix="oosit_docx_")
        
        # Get strategy names from config
        default_strategies = self.config.get('default_strategies', [])
        test_strategies = self.config.get('testing_strategies', [])
        all_strategies = default_strategies + test_strategies
        
        if not all_strategies:
            logger.error("No strategies found in config")
            return
        
        # Process all strategies
        # Note: keeping variable name for compatibility
        all_strategies = all_strategies
        
        for strategy_name in all_strategies:
            try:
                # Search for the strategy .py file in main directory and saved subfolder
                py_file = None
                search_directories = [
                    self.strategies_directory,  # Main strategies directory
                    self.strategies_directory / "saved"  # Saved subfolder
                ]
                
                for search_dir in search_directories:
                    if not search_dir.exists():
                        continue
                    candidate_file = search_dir / f"{strategy_name}.py"
                    if candidate_file.exists():
                        py_file = candidate_file
                        break
                
                if py_file is None:
                    logger.error(f"Strategy file not found: {strategy_name}.py (searched in {[str(d) for d in search_directories]})")
                    continue
                
                # Import the module to get _explanation
                spec = importlib.util.spec_from_file_location(strategy_name, py_file)
                if spec is None or spec.loader is None:
                    logger.error(f"Could not create spec for {py_file}")
                    continue
                
                module = importlib.util.module_from_spec(spec)
                spec.loader.exec_module(module)
                
                # Get explanation from module
                explanation = getattr(module, '_explanation', '')
                
                # Create DOCX file with proper structure from blah.txt
                doc = Document()
                
                # Get font settings from config
                font_name = self.config.get('font_name', '바탕')
                font_size = self.config.get('font_size', 11)
                
                # Add strategy name
                p = doc.add_paragraph(strategy_name)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in p.runs:
                    run.font.name = font_name
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = Pt(font_size)
                
                # Add creation date placeholder
                p = doc.add_paragraph('작성일자: ')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                
                # Add section 1
                p = doc.add_paragraph('1. 전략의 설명')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in p.runs:
                    run.font.name = font_name
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = Pt(font_size)
                
                # Add explanation
                p = doc.add_paragraph(explanation.strip())
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in p.runs:
                    run.font.name = font_name
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = Pt(font_size)
                
                # Add section 2
                p = doc.add_paragraph('2. 전략의 분석')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in p.runs:
                    run.font.name = font_name
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = Pt(font_size)
                
                # Save the document
                docx_path = Path(self.temp_docx_dir) / f"{strategy_name}_description.docx"
                doc.save(str(docx_path))
                logger.debug(f"Created DOCX for strategy: {strategy_name}")
                
            except Exception as e:
                logger.error(f"Error creating DOCX for {strategy_name}: {e}")
    
    def _generate_strategy_reports(self, backtest_results, 
                                 config):
        """Generate individual Word document reports for each strategy."""
        results = backtest_results['results']
        summary = backtest_results['summary']
        periods = backtest_results['periods']
        
        default_strategy_names = summary['default_strategy_names']
        
        if not self.temp_docx_dir:
            logger.warning("No DOCX files extracted, skipping strategy reports")
            return
        
        docx_files = list(Path(self.temp_docx_dir).glob('*.docx'))
        
        for docx_file in docx_files:
            # Determine corresponding strategy
            strategy_name = docx_file.stem.replace('_description', '')
            
            if strategy_name not in results:
                logger.warning(f"No results found for strategy corresponding to {docx_file.name}")
                continue
            
            try:
                self._generate_single_strategy_report(
                    docx_file, strategy_name, results[strategy_name],
                    default_strategy_names, periods, config
                )
                logger.info(f"Generated report for {strategy_name}")
                
            except Exception as e:
                logger.error(f"Error generating report for {strategy_name}: {e}")
    
    def _generate_single_strategy_report(self, docx_file, strategy_name,
                                       strategy_results,
                                       default_strategy_names,
                                       periods, config):
        """Generate a single strategy report."""
        # Update creation date in document
        doc = Document(docx_file)
        today_date = datetime.today().strftime("%Y.%m.%d")
        
        # Get font settings from config
        font_name = config.get('font_name', '바탕')
        font_size = config.get('font_size', 11)
        
        for paragraph in doc.paragraphs:
            if "작성일자" in paragraph.text or "작성 일자" in paragraph.text:
                paragraph.text += today_date
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                # Apply font styling from config
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = Pt(font_size)
                break
        
        doc.save(docx_file)
        
        # Set up document logger with config font settings
        logger_doc = DocxLogger(str(docx_file))
        font_name = config.get('font_name', '바탕')
        font_size = config.get('font_size', 11)
        logger_doc.set_default_style(font_name=font_name, font_size=font_size, space_after=0)
        
        # Redirect stdout to document
        original_stdout = sys.stdout
        original_show = plt.show
        
        def custom_show(*args, **kwargs):
            fig = plt.gcf()
            logger_doc.add_plot(fig)
            logger_doc.flush()
            plt.close(fig)
        
        try:
            sys.stdout = logger_doc
            plt.show = custom_show
            
            # Write report content
            self._write_strategy_report_content(
                strategy_name, strategy_results, default_strategy_names, periods, config
            )
            
            logger_doc.flush()
            
        finally:
            sys.stdout = original_stdout
            plt.show = original_show
        
        # Move file to docx results directory with new name
        docx_dir = self.docx_dir
        base_name = docx_file.stem.replace('_description', '')
        
        # Add * prefix for default strategies
        default_strategies = self.config.get('default_strategies', [])
        if base_name in default_strategies:
            base_name = f"_{base_name}"
        
        # Determine additional flag for filename
        additional_flag = ''
        config_file = self.config.get('config_file', None)
        if config_file:
            config_path = Path(config_file)
            config_name = config_path.stem  # Get filename without extension
            
            if config_name == 'default_config':
                additional_flag = ''
            elif config_name == 'dxy_config':
                additional_flag = ' [DXY]'
            elif config_name == 'veu_redirect_config':
                additional_flag = ' [VEU]'
            else:
                additional_flag = f' [{config_name}]'
        
        new_name = docx_dir / f"{base_name}{additional_flag}.docx"
        docx_file.rename(new_name)
    
    def _write_strategy_report_content(self, strategy_name,
                                     strategy_results,
                                     default_strategy_names,
                                     periods, config):
        """Write the main content of a strategy report."""
        # Determine if this is a test strategy from config
        default_strategies = self.config.get('default_strategies', [])
        is_test_strategy = strategy_name not in default_strategies
        
        # Header information
        print('- 총 기간(및 닷컴버블 하락장)을 제외한 다른 분석은 QQQ의 가격변동이 없도록 하는 하락장 기간들에 대한 분석임.')
        print('- 그래프에서 y축은 포트폴리오의 가치를 % 단위로 표현한 것임. 가치는 100%에서 시작함. 또한 log scale이 적용됨.')
        print(f"- 사용한 대조군 전략들: {', '.join(default_strategy_names)}")
        print('- 괄호 안의 수치들은 위 대조군 전략들의 수치를 차례로 나열한 것임.')
        print()  # Empty line
        
        # Full period analysis
        full_period_result = strategy_results.get("Full Period")
        if full_period_result:
            print(f'총 분석 기간 : {config["full_start_date"]} - {config["full_end_date"]}')
            # Get default strategy results for this period
            default_results = self._get_default_strategy_results("Full Period")
            self._plot_comparison(full_period_result, default_results, is_test_strategy)
            print()
        
        # Period-specific analyses
        for i, period_config in enumerate(config.get('test_periods', [])):
            period_name = period_config['period_name']
            if period_name in strategy_results:
                result = strategy_results[period_name]
                print(f'{i+1}) {period_name} : {period_config["period_start_date"]} - {period_config["period_end_date"]}')
                # Get default strategy results for this period
                default_results = self._get_default_strategy_results(period_name)
                self._plot_comparison(result, default_results, is_test_strategy)
                print()
    
    def _get_default_strategy_results(self, period_name):
        """
        Get default strategy results for a specific period.
        
        Args:
            period_name: Name of the period to get results for
            
        Returns:
            List of BacktestResult objects for default strategies
        """
        if not self.backtest_results:
            return []
        
        # Get default strategy names from config
        default_strategy_names = self.config.get('default_strategies', [])
        
        if not default_strategy_names:
            logger.debug("No default strategies found in config")
            return []
        
        results = self.backtest_results['results']
        default_results = []
        
        for strategy_name, strategy_results in results.items():
            # Check if this strategy is in the default strategies list
            if strategy_name in default_strategy_names:
                if period_name in strategy_results:
                    default_results.append(strategy_results[period_name])
        
        return default_results
    
    def _plot_comparison(self, result, default_results,
                        is_test):
        """Plot comparison chart and print summary statistics."""
        plt.figure(figsize=(10, 6))
        
        # Plot test strategy if it's a test strategy
        if is_test:
            plt.plot(result.date_range, result.normalized_values, 
                    label=result.display_name, linewidth=2)
        
        # Plot default strategies (restored from original implementation)
        for default_result in default_results:
            plt.plot(default_result.date_range, default_result.normalized_values, 
                    label=default_result.display_name, linewidth=2)
        
        plt.xlabel('Date')
        plt.ylabel('Portfolio Value in %, log scale')
        plt.yscale('log')
        plt.legend(bbox_to_anchor=(1.02, 1), loc='upper left', borderaxespad=0, fontsize=12)
        plt.tight_layout()
        plt.show()
        
        # Print summary statistics with default comparison (restored from original)
        default_returns = [dr.total_return for dr in default_results]
        default_drawdowns = [dr.max_drawdown for dr in default_results]
        
        # Format default returns and drawdowns for display
        if default_returns:
            default_returns_str = "(" + ", ".join([f"{r:.1f}%" for r in default_returns]) + ")"
            default_drawdowns_str = "(" + ", ".join([f"-{d:.1f}%" for d in default_drawdowns]) + ")"
        else:
            default_returns_str = ""
            default_drawdowns_str = ""
        
        print(f'수익률 : {result.total_return:.1f}% {default_returns_str}')
        print(f'최대 낙폭 : -{result.max_drawdown:.1f}% {default_drawdowns_str}')
    
    def _generate_csv_summaries(self, backtest_results, 
                              config):
        """Generate CSV summary files."""
        summary = backtest_results['summary']
        
        # Get results directory (already created and cleaned in _setup_result_directories)
        csv_dir = self.csv_dir
        
        try:
            # Save summary DataFrames
            summary['full_result_df'].to_csv(
                csv_dir / "총 기간 분석.csv", 
                index=False, encoding="utf-8-sig"
            )
            
            summary['periods_return_df'].to_csv(
                csv_dir / "기간별 회복 수익률 (%).csv", 
                index=False, encoding="utf-8-sig"
            )
            
            summary['periods_maxdd_df'].to_csv(
                csv_dir / "기간별 최대 낙폭 (%).csv", 
                index=False, encoding="utf-8-sig"
            )
            
            summary['rebalancing_log_df'].to_csv(
                csv_dir / "리밸런싱 로그.csv", 
                index=False, encoding="utf-8-sig"
            )
            
            logger.info("Generated CSV summary files")
            
        except Exception as e:
            logger.error(f"Error generating CSV summaries: {e}")
            raise
    
    def _generate_archive(self, config):
        """Generate TGZ archive with configuration and scripts."""
        # Generate archive name based on test strategies and config file
        test_strategies = config.get('testing_strategies', [])
        config_file = config.get('config_file', None)
        
        # Create comma-joined string of test strategies
        strategies_str = ','.join(test_strategies) if test_strategies else 'no_test_strategies'
        
        # Determine additional flag
        additional_flag = ''
        if config_file:
            config_path = Path(config_file)
            config_name = config_path.stem  # Get filename without extension
            
            if config_name == 'default_config':
                additional_flag = ''
            elif config_name == 'dxy_config':
                additional_flag = ' [DXY]'
            elif config_name == 'veu_redirect_config':
                additional_flag = ' [VEU]'
            else:
                additional_flag = f' [{config_name}]'
        
        # Create archive name
        archive_name = f"{strategies_str}{additional_flag}.tgz"
        tgz_path = self.timestamp_dir / archive_name
        temp_json_path = 'configure.json'
        
        try:
            # Write configuration to temporary JSON file
            with open(Path(temp_json_path), 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=4)
            
            # Create TGZ archive
            with tarfile.open(Path(tgz_path), "w:gz") as tar:
                # Add configuration file
                tar.add(temp_json_path, arcname='configure.json')
                
                # Get strategy names from config
                default_strategies = config.get('default_strategies', [])
                test_strategies = config.get('testing_strategies', [])
                all_strategies = default_strategies + test_strategies
                
                # Check if target.json exists to add to archive
                target_json_path = Path('./jsons/target.json')
                if all_strategies:
                    
                    # Add only the Python files specified in target.json
                    search_directories = [
                        self.strategies_directory,  # Main strategies directory
                        self.strategies_directory / "saved"  # Saved subfolder
                    ]
                    
                    for strategy_name in all_strategies:
                        strategy_added = False
                        for search_dir in search_directories:
                            if not search_dir.exists():
                                continue
                            py_file = search_dir / f"{strategy_name}.py"
                            if py_file.exists():
                                try:
                                    # Add .py files directly to archive root (not in subfolders)
                                    tar.add(py_file, arcname=py_file.name)
                                    logger.debug(f"Added {py_file.name} to archive")
                                    strategy_added = True
                                    break
                                except Exception as e:
                                    logger.warning(f"Error adding {py_file} to archive: {e}")
                        
                        if not strategy_added:
                            logger.warning(f"Strategy file not found: {strategy_name}.py")
                    
                    # Add target.json itself if it exists
                    if target_json_path.exists():
                        tar.add(target_json_path, arcname='target.json')
                        logger.debug("Added target.json to archive")
                    else:
                        logger.warning("target.json not found - not added to archive")
                else:
                    logger.warning("No strategies found in config")
            
            logger.info("Generated TGZ archive")
            
            # Extract archive to run_archive directory
            with tarfile.open(Path(tgz_path), "r:gz") as tar:
                tar.extractall(self.archive_dir)
                logger.info(f"Extracted archive to {self.archive_dir}")
            
        except Exception as e:
            logger.error(f"Error generating archive: {e}")
            raise
        finally:
            # Clean up temporary JSON file
            try:
                os.remove(temp_json_path)
            except:
                pass
    
    def _cleanup_temp_directories(self):
        """Clean up temporary directories."""
        if self.temp_docx_dir and Path(self.temp_docx_dir).exists():
            shutil.rmtree(self.temp_docx_dir)
            logger.debug(f"Cleaned up temporary DOCX directory: {self.temp_docx_dir}")
    
    def compare_strategies(self, strategy_result, 
                         default_results):
        """
        Compare a strategy against default strategies.
        
        Args:
            strategy_result: Result for the strategy to compare
            default_results: List of default strategy results
            
        Returns:
            Comparison metrics and data
        """
        comparison = {
            'strategy_name': strategy_result.display_name,
            'strategy_return': strategy_result.total_return,
            'strategy_drawdown': strategy_result.max_drawdown,
            'default_returns': [r.total_return for r in default_results],
            'default_drawdowns': [r.max_drawdown for r in default_results],
            'default_names': [r.display_name for r in default_results]
        }
        
        return comparison
    
    def create_performance_chart(self, results, 
                               title="Performance Comparison"):
        """
        Create a performance comparison chart.
        
        Args:
            results: List of BacktestResult objects to plot
            title: Chart title
            
        Returns:
            Matplotlib figure
        """
        fig, ax = plt.subplots(figsize=(12, 8))
        
        for result in results:
            ax.plot(result.date_range, result.normalized_values, 
                   label=result.display_name, linewidth=2)
        
        ax.set_xlabel('Date')
        ax.set_ylabel('Portfolio Value (%)')
        ax.set_yscale('log')
        ax.set_title(title)
        ax.legend(bbox_to_anchor=(1.05, 1), loc='upper left', fontsize=12)
        ax.grid(True, alpha=0.3)
        
        plt.tight_layout()
        return fig